import json
import os
import re
import docx

QUESTIONS_PATH = r'C:\Users\丁泽伟\Downloads\课后习题及答案\课后习题及答案\《Spark大数据分析与实战（第2版）》课后习题.docx'
ANSWERS_PATH = r'C:\Users\丁泽伟\Downloads\课后习题及答案\课后习题及答案\《Spark大数据分析与实战（第2版）》课后习题答案.docx'
OUTPUT_PATH = r'D:\mimo\quiz-app\public\spark-questions.json'

FILL_SECTION = re.compile(r'^[一二三四五六七八九十]*[、．.]?\s*填空题')
BOOL_SECTION = re.compile(r'^[一二三四五六七八九十]*[、．.]?\s*判断题')
CHOICE_SECTION = re.compile(r'^[一二三四五六七八九十]*[、．.]?\s*选择题')
ESSAY_SECTION = re.compile(r'^[一二三四五六七八九十]*[、．.]?\s*(简答题|编程题)')
NUMBERED = re.compile(r'^(\d+)[．.、\s]')
LETTER_OPTION = re.compile(r'^([A-Z])[.．、\s]')
HAS_BLANK = re.compile(r'[　\s]{2,}|_+')
HAS_BOOL_PAREN = re.compile(r'[（(]\s{2,}[)）]|[（(]\s*[)）]\s*$')
MULTI_MARKER = re.compile(r'[（(]\s*多选\s*[)）]')


def is_chapter_style(style_name):
    return '节' in style_name


def parse_questions_doc(doc):
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append((text, para.style.name))

    chapters = []
    current_chapter = None
    current_type = None
    q_counter = {}

    i = 0
    while i < len(paragraphs):
        text, style = paragraphs[i]

        if is_chapter_style(style):
            current_chapter = {"id": f"ch{len(chapters)+1}", "name": text, "questions": []}
            chapters.append(current_chapter)
            current_type = None
            q_counter = {}
            i += 1
            continue

        if not current_chapter:
            i += 1
            continue

        if FILL_SECTION.match(text):
            current_type = "fill"; q_counter.setdefault("fill", 0); current_stem = None; i += 1; continue
        if BOOL_SECTION.match(text):
            current_type = "boolean"; q_counter.setdefault("boolean", 0); current_stem = None; i += 1; continue
        if CHOICE_SECTION.match(text):
            current_type = "choice"; q_counter.setdefault("choice", 0); current_stem = None; i += 1; continue
        if ESSAY_SECTION.match(text):
            m = ESSAY_SECTION.match(text)
            current_type = "code" if '编程' in m.group(1) else "essay"
            q_counter.setdefault(current_type, 0)
            current_stem = None
            i += 1
            continue

        if not current_type:
            i += 1
            continue

        if current_type == "fill":
            if HAS_BLANK.search(text):
                q_counter["fill"] = q_counter.get("fill", 0) + 1
                current_chapter["questions"].append({
                    "id": f"{current_chapter['id']}-fill-{q_counter['fill']}",
                    "type": "fill", "stem": text, "answer": [], "explanation": ""
                })
            i += 1
            continue

        if current_type == "boolean":
            if HAS_BOOL_PAREN.search(text):
                q_counter["boolean"] = q_counter.get("boolean", 0) + 1
                stem = re.sub(r'[（(]\s*[)）]\s*$', '', text).strip()
                current_chapter["questions"].append({
                    "id": f"{current_chapter['id']}-bool-{q_counter['boolean']}",
                    "type": "boolean", "stem": stem, "answer": None, "explanation": ""
                })
            i += 1
            continue

        if current_type == "choice":
            if HAS_BOOL_PAREN.search(text):
                q_counter["choice"] = q_counter.get("choice", 0) + 1
                stem = text
                is_multi = bool(MULTI_MARKER.search(text))
                options = []

                stem_parts = re.split(r'\t+|\s{3,}', text)
                if len(stem_parts) > 1:
                    stem = stem_parts[0].strip()
                    for part in stem_parts[1:]:
                        part = part.strip()
                        om = LETTER_OPTION.match(part)
                        if om:
                            options.append({"label": om.group(1), "text": re.sub(r'^[A-Z][.．、\s]+', '', part)})

                j = i + 1
                while j < len(paragraphs):
                    nt, ns = paragraphs[j]
                    if is_chapter_style(ns) or FILL_SECTION.match(nt) or BOOL_SECTION.match(nt) or \
                       CHOICE_SECTION.match(nt) or ESSAY_SECTION.match(nt) or \
                       HAS_BOOL_PAREN.search(nt):
                        break
                    opt_m = LETTER_OPTION.match(nt)
                    if opt_m:
                        options.append({"label": opt_m.group(1), "text": re.sub(r'^[A-Z][.．、\s]+', '', nt)})
                        j += 1
                    else:
                        parts = re.split(r'\t+|\s{3,}', nt)
                        any_opt = False
                        for part in parts:
                            part = part.strip()
                            om = LETTER_OPTION.match(part)
                            if om:
                                options.append({"label": om.group(1), "text": re.sub(r'^[A-Z][.．、\s]+', '', part)})
                                any_opt = True
                        if any_opt:
                            j += 1
                        else:
                            break
                current_chapter["questions"].append({
                    "id": f"{current_chapter['id']}-choice-{q_counter['choice']}",
                    "type": "multiple" if is_multi else "single",
                    "stem": stem, "options": options, "answer": [], "explanation": ""
                })
                i = j
                continue
            i += 1
            continue

        if current_type in ("essay", "code"):
            m = NUMBERED.match(text)
            if m:
                q_counter[current_type] = q_counter.get(current_type, 0) + 1
                stem = re.sub(r'^\d+[．.、\s]+', '', text)
                current_chapter["questions"].append({
                    "id": f"{current_chapter['id']}-{current_type}-{q_counter[current_type]}",
                    "type": current_type, "stem": stem, "answer": "", "explanation": ""
                })
            else:
                if current_chapter["questions"] and current_chapter["questions"][-1]["type"] == current_type:
                    current_chapter["questions"][-1]["stem"] += "\n" + text
                else:
                    q_counter[current_type] = q_counter.get(current_type, 0) + 1
                    current_chapter["questions"].append({
                        "id": f"{current_chapter['id']}-{current_type}-{q_counter[current_type]}",
                        "type": current_type, "stem": text, "answer": "", "explanation": ""
                    })
            i += 1
            continue

        i += 1

    return chapters


def parse_answers_doc(doc):
    answers = {}
    current_chapter_idx = -1
    current_type = None
    question_counter = 0

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append((text, para.style.name))

    i = 0
    while i < len(paragraphs):
        text, style = paragraphs[i]

        if is_chapter_style(style):
            current_chapter_idx += 1
            current_type = None
            question_counter = 0
            i += 1
            continue

        if '例程代码' in style:
            if current_type in ("essay", "code") and current_chapter_idx >= 0 and question_counter > 0:
                key = f"{current_chapter_idx}-{current_type}-{question_counter}"
                if key not in answers:
                    answers[key] = ""
                answers[key] += text + "\n"
            i += 1
            continue

        if FILL_SECTION.match(text): current_type = "fill"; question_counter = 0; i += 1; continue
        if BOOL_SECTION.match(text): current_type = "boolean"; question_counter = 0; i += 1; continue
        if CHOICE_SECTION.match(text): current_type = "choice"; question_counter = 0; i += 1; continue
        if ESSAY_SECTION.match(text):
            m = ESSAY_SECTION.match(text)
            current_type = "code" if '编程' in m.group(1) else "essay"
            question_counter = 0
            i += 1
            continue

        if not current_type:
            i += 1
            continue

        if current_type in ("fill", "boolean", "choice"):
            m = NUMBERED.match(text)
            if m:
                question_counter += 1
                answer_text = re.sub(r'^\d+[．.、\s]+', '', text).strip()
                key = f"{current_chapter_idx}-{current_type}-{question_counter}"
                answers[key] = answer_text
            i += 1
            continue

        if current_type in ("essay", "code"):
            m = NUMBERED.match(text)
            if m:
                question_counter += 1
                key = f"{current_chapter_idx}-{current_type}-{question_counter}"
                answers[key] = ""
                i += 1
                continue
            if question_counter > 0:
                key = f"{current_chapter_idx}-{current_type}-{question_counter}"
                if key in answers:
                    if answers[key]:
                        answers[key] += "\n" + text
                    else:
                        answers[key] = text
            i += 1
            continue

        i += 1

    return answers


def merge_answers(chapters, answers):
    answer_idx = {}
    for key, val in answers.items():
        parts = key.split('-')
        ch_idx = int(parts[0])
        q_type = parts[1]
        if ch_idx not in answer_idx:
            answer_idx[ch_idx] = {}
        if q_type not in answer_idx[ch_idx]:
            answer_idx[ch_idx][q_type] = []
        answer_idx[ch_idx][q_type].append(val)

    type_counters = {}
    for ch_i, chapter in enumerate(chapters):
        type_counters[ch_i] = {}
        for q in chapter["questions"]:
            q_type = q["type"]
            lookup_type = "choice" if q_type in ("single", "multiple") else q_type
            if q_type not in type_counters[ch_i]:
                type_counters[ch_i][q_type] = 0
            type_counters[ch_i][q_type] += 1
            counter = type_counters[ch_i][q_type]

            ans_list = answer_idx.get(ch_i, {}).get(lookup_type, [])
            if counter <= len(ans_list):
                raw_answer = ans_list[counter - 1]
                apply_answer(q, raw_answer)


def apply_answer(question, raw_answer):
    q_type = question["type"]

    if q_type == "fill":
        answers = re.split(r'[、，,;；]', raw_answer)
        question["answer"] = [a.strip() for a in answers if a.strip()]

    elif q_type == "boolean":
        question["answer"] = '对' in raw_answer or '√' in raw_answer

    elif q_type in ("single", "multiple"):
        letters = re.findall(r'[A-Z]', raw_answer)
        if q_type == "single" and len(letters) == 1:
            question["answer"] = letters[0]
        else:
            question["answer"] = letters

    elif q_type in ("essay", "code"):
        question["answer"] = raw_answer.strip()


def fix_choice_types(chapters):
    for chapter in chapters:
        for q in chapter["questions"]:
            if q["type"] in ("single", "multiple"):
                ans = q.get("answer", [])
                if isinstance(ans, list):
                    if len(ans) > 1:
                        q["type"] = "multiple"
                    elif len(ans) == 1:
                        q["type"] = "single"
                        q["answer"] = ans[0]


def main():
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

    print("Parsing questions...")
    q_doc = docx.Document(QUESTIONS_PATH)
    chapters = parse_questions_doc(q_doc)
    print(f"Found {len(chapters)} chapters")
    for ch in chapters:
        print(f"  {ch['name']}: {len(ch['questions'])} questions")

    print("\nParsing answers...")
    a_doc = docx.Document(ANSWERS_PATH)
    answers = parse_answers_doc(a_doc)
    print(f"Found {len(answers)} answer entries")

    print("\nMerging...")
    merge_answers(chapters, answers)
    fix_choice_types(chapters)

    for ch in chapters:
        for q in ch["questions"]:
            if q["type"] in ("essay", "code") and not q["answer"]:
                q["answer"] = "（参考答案暂缺，请参阅教材相关内容）"

    total = sum(len(ch["questions"]) for ch in chapters)
    data = {
        "meta": {
            "subject": "Spark大数据分析与实战（第2版）",
            "version": "1.0",
            "totalQuestions": total
        },
        "chapters": chapters
    }

    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"\nOutput: {OUTPUT_PATH}")
    print(f"Total questions: {total}")

    no_answer = 0
    for ch in chapters:
        for q in ch["questions"]:
            if q["type"] == "fill" and not q["answer"]:
                no_answer += 1
            elif q["type"] == "boolean" and q["answer"] is None:
                no_answer += 1
            elif q["type"] in ("single", "multiple") and not q["answer"]:
                no_answer += 1
            elif q["type"] in ("essay", "code") and not q["answer"]:
                no_answer += 1
    if no_answer:
        print(f"WARNING: {no_answer} questions without answers!")
    else:
        print("All questions have answers!")


if __name__ == "__main__":
    main()
